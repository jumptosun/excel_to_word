#!/usr/bin/env python3

from dataclasses import replace
from re import template
import sys, os
import argparse
import pandas as pd
from docx import Document 
import copy

class ExcelParser:
    def __init__(self):
        self.df = pd.DataFrame()
        self.header_patterns = []

    def  parse_excel(self, filename, sheetname):
        self.df = pd.read_excel(filename, sheet_name = sheetname)
        self.__parse_header_pattern()

    def __parse_header_pattern(self):
        header = self.df.head()
        for v in header:
            l = v.split("/")
            if len(l) > 1:
                self.header_patterns.append(l[-1])
            else:
                self.header_patterns.append("")


class WordPatternReplacer:
    def __init__(self):
        return

    def read_word(self, wordfile):
        self.__word_template_orign = Document(wordfile)
        self.__word_template = copy.deepcopy(self.__word_template_orign)

    def reset(self):
        del(self.__word_template)
        self.__word_template = copy.deepcopy(self.__word_template_orign)

    def replace_pattern(self, pattern, target):
        for p in self.__word_template.paragraphs:
            for r in p.runs:
                r.text = r.text.replace(pattern, target)

    # patterns, targets should be iterable
    def replace_patterns(self, patterns, targets):
        assert len(patterns) == len(targets)

        for p in self.__word_template.paragraphs:
            for r in p.runs:
                for i in range(len(patterns)):
                    if patterns[i] != "":
                        r.text = r.text.replace(patterns[i], str(targets[i]))

    def save_file(self, filename):
        self.__word_template.save(filename)

def replace_pattern_in_template(excel_filepath: str, sheet_name: str, template_word_filepath: str):
    # parse xls
    ep = ExcelParser()
    ep.parse_excel(excel_filepath, sheet_name)

    # read template
    wp = WordPatternReplacer()
    wp.read_word(template_word_filepath)

    template_word_filename = os.path.basename(template_word_filepath)

    for i in range(len(ep.df)):
        wp.reset()
        wp.replace_patterns(ep.header_patterns, ep.df.iloc[i,:].to_list())
        wp.save_file("%s-%s"%(ep.df.iloc[i, 0], template_word_filename))


if __name__ == "__main__":
    # examine the parameter
    parser = argparse.ArgumentParser()
    parser.add_argument("-x", nargs = 2, dest = "excel_file", help="input excel file")
    parser.add_argument("-w", dest = "template_word_file", help="input temple word file")
    args = parser.parse_args()

    if not args.excel_file or not args.template_word_file:
        parser.print_help()
        exit(0)
    else:
        print("""the input file:
        source excel file: %s
        template word file: %s
        """ % (args.excel_file, args.template_word_file))

    replace_pattern_in_template(args.excel_file[0], args.excel_file[1], args.template_word_file)
